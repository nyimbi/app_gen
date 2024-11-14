import sqlalchemy as sa
from sqlalchemy import inspect, Boolean, Date, DateTime, Enum, Float, Integer, Numeric, String, Text, Time, ForeignKey, Table, ForeignKeyConstraint, PrimaryKeyConstraint, MetaData, create_engine
from sqlalchemy.orm import RelationshipProperty
from sqlalchemy.sql import sqltypes
import inflect
import math
import argparse
from typing import Any, Dict, List, Optional, Union
from flask import g, flash, redirect, url_for, session, request, render_template, make_response
from flask_appbuilder.models.sqla.interface import SQLAInterface
from flask_appbuilder import ModelView, MasterDetailView, MultipleView, ModelRestApi
from flask_appbuilder.fields import AJAXSelectField, QuerySelectField, QuerySelectMultipleField, EnumField
from flask_appbuilder.fieldwidgets import Select2AJAXWidget, Select2SlaveAJAXWidget, Select2ManyWidget, BS3TextFieldWidget, BS3PasswordFieldWidget, DatePickerWidget, DateTimePickerWidget, Select2Widget
from flask_appbuilder.actions import action
from flask_appbuilder.security.decorators import has_access
from flask_appbuilder.forms import DynamicForm
from flask_appbuilder import AppBuilder, BaseView, expose, has_access
from flask_appbuilder.charts.views import GroupByChartView, ChartView
from flask_appbuilder.models.group import aggregate_count
from flask_appbuilder.widgets import ListThumbnail, FormVerticalWidget
from flask_login import current_user
from wtforms import StringField, BooleanField, IntegerField, FloatField, DateField, DateTimeField, SelectField, HiddenField, TextAreaField, DecimalField, validators
from cryptography.fernet import Fernet
import PyPDF2
import pypandoc
import magic
import docx
import nltk
from nltk.corpus import stopwords
from nltk.tokenize import sent_tokenize, word_tokenize
from transformers import pipeline
import markdown
import pdfkit
from io import BytesIO
import mimetypes
import json
import graphene
from graphene_sqlalchemy import SQLAlchemyObjectType, SQLAlchemyConnectionField

from view_utils import get_view_icon
from utils import snake_to_pascal, snake_to_words, pascal_to_words, get_class_name

# Set up logging
import logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Initialize inflect engine
p = inflect.engine()

# Global variables
INDENT = "    "
AB_PREFIX = 'ab_'
generated_views = []
generated_code = []

def generate_views(db_uri: str) -> None:
    """
    Generate Flask-AppBuilder views for all tables in the database.

    This function is the main entry point for view generation. It connects to the database,
    reflects the schema, and generates various types of views including ModelViews,
    MasterDetailViews, MultipleViews, WizardViews, ChartViews, and API views.

    Args:
        db_uri (str): SQLAlchemy database URI to connect to

    Returns:
        None
    """
    engine = sa.create_engine(db_uri)
    inspector = inspect(engine)
    metadata = sa.MetaData()
    metadata.reflect(bind=engine)

    # Add imports and initial setup
    generated_code.append([
        "imports",
        """from flask_appbuilder import ModelView, MasterDetailView, MultipleView, ModelRestApi
from flask_appbuilder.models.sqla.interface import SQLAInterface
from flask_appbuilder.fieldwidgets import BS3TextFieldWidget, Select2Widget, DatePickerWidget
from flask_appbuilder.forms import DynamicForm
from wtforms import StringField, IntegerField, DateField, SelectField
from flask_appbuilder.actions import action
from flask import session, redirect, url_for, flash, Markup
from flask_appbuilder.baseviews import expose, BaseView
from flask_appbuilder.widgets import FormVerticalWidget, ListThumbnail
from flask_appbuilder.charts.views import GroupByChartView
from flask_appbuilder.models.group import aggregate_count
from flask_appbuilder.api import BaseApi, expose
from flask import flash, redirect, request, render_template, make_response, url_for
from flask_appbuilder.security.decorators import has_access
from flask_login import current_user
from . import appbuilder, db
from .models import *
from .view_utils import get_view_icon
import datetime
import graphene
from graphene_sqlalchemy import SQLAlchemyObjectType, SQLAlchemyConnectionField

# Global list to store all generated views
generated_views = []

class BeautifulListWidget(ListThumbnail):
    template = 'beautiful_list.html'

class BeautifulFormWidget(FormVerticalWidget):
    template = 'beautiful_form.html'

class WizardView(BaseView):
    default_view = 'step1'

    def _init_steps(self):
        self.steps = {}
        for step in range(1, self.total_steps + 1):
            step_name = f'step{step}'
            self.steps[step_name] = getattr(self, step_name)

    def is_step_complete(self, step):
        return session.get(f'{self.__class__.__name__}_step{step}_complete', False)

    def mark_step_complete(self, step):
        session[f'{self.__class__.__name__}_step{step}_complete'] = True

    def get_progress(self):
        completed_steps = sum(self.is_step_complete(step) for step in range(1, self.total_steps + 1))
        return (completed_steps / self.total_steps) * 100

    @expose('/reset')
    def reset(self):
        for step in range(1, self.total_steps + 1):
            session.pop(f'{self.__class__.__name__}_step{step}_complete', None)
        flash('Your progress has been reset. Start again from the beginning!', 'info')
        return redirect(url_for(f'.step1'))

    def render_wizard(self, step, form, step_description):
        return self.render_template(
            'wizard.html',
            form=form,
            progress=self.get_progress(),
            current_step=step,
            total_steps=self.total_steps,
            step_description=step_description
        )

class SessionManagedView(BaseView):
    @expose("/")
    def list(self):
        return self.render_template("list.html", partial_data=session.get("partial_form_data"))

    @expose("/add", methods=["GET", "POST"])
    def add(self):
        form = self.add_form()
        if form.validate_on_submit():
            self._save_form_to_db(form)
            session.pop("partial_form_data", None)
            return redirect(url_for(f"{self.__class__.__name__}.list"))
        else:
            session["partial_form_data"] = form.data
        return self.render_template("add.html", form=form)
"""
    ])

    for table_name in inspector.get_table_names():
        table = metadata.tables[table_name]
        generate_model_view(table, p)
        generate_master_detail_views(table, metadata, p)
        generate_multiple_views(table, metadata, p)
        generate_wizard_view(table, p)
        generate_graphql(table, p)
        generate_ModelRestApi(table, p)
        generate_chart_view(table, p)

    # Add view registration function
    generated_code.append([
        "register_views",
        """
def register_views():
""" + "\n".join([f"    appbuilder.add_view({class_name}, '{table_name}', icon='{icon}', category='Views')"
                  for class_name, view_type, table_name, icon in generated_views]) + """

    # Gamification: Reminder for incomplete forms
    @appbuilder.app.context_processor
    def inject_notifications():
        notifications = []
        for view_name, view in appbuilder.baseviews.items():
            if isinstance(view, WizardView):
                progress = view.get_progress()
                if 0 < progress < 100:
                    notifications.append(f"Continue your {view.__class__.__name__[:-10]} form! You're {progress:.0f}% done.")
        return dict(notifications=notifications)

# GraphQL schema
schema = graphene.Schema(query=Query)
"""
    ])

def get_list_columns(columns: List[sa.Column]) -> List[str]:
    """
    Get a list of column names suitable for list views, excluding certain system columns.

    Args:
        columns (List[sa.Column]): List of SQLAlchemy Column objects

    Returns:
        List[str]: List of column names suitable for list views
    """
    ignored_fields = {'id', 'updated_at', 'updated_by', 'created_at', 'created_by', 'is_bookmarked', 'is_archived'}
    return [col.name for col in columns if col.name not in ignored_fields]

def get_field_sets(table: sa.Table) -> Dict[str, List[Tuple[str, str]]]:
    """
    Generate field sets for forms, grouping columns into sets of 4.

    Args:
        table (sa.Table): SQLAlchemy Table object

    Returns:
        Dict[str, List[Tuple[str, str]]]: Dictionary of field sets with labels
    """
    ignored_fields = {'id', 'updated_at', 'updated_by', 'created_at', 'created_by', 'is_bookmarked', 'is_archived'}
    valid_columns = [col for col in table.columns if col.name not in ignored_fields]

    # Group fields into sets of 4
    field_sets = [valid_columns[i:i+4] for i in range(0, len(valid_columns), 4)]

    # Create field_sets dictionary
    field_sets_dict = {
        f"Field Set {i+1}": [
            (col.name, col.name.replace('_', ' ').title()) for col in field_set
        ] for i, field_set in enumerate(field_sets)
    }

    return field_sets_dict

def generate_model_view(table: sa.Table, p: inflect.engine) -> None:
    """
    Generate a comprehensive ModelView for a table.

    This function creates a ModelView class for the given table, including list columns,
    form configurations, actions, and other view-specific settings.

    Args:
        table (sa.Table): SQLAlchemy Table object
        p (inflect.engine): Inflect engine for plural/singular conversions

    Returns:
        None
    """
    class_name = get_class_name(table.name, p)
    view_name = f"{class_name}ModelView"
    columns = [c.name for c in table.columns]
    list_columns = get_list_columns(columns)
    field_sets = get_field_sets(table)
    icon = get_view_icon(table.name, "ModelView")

    view_code = [
        f"class {view_name}(ModelView):",
        f"    datamodel = SQLAInterface({class_name})",
        f"    list_columns = {list_columns}",
        f"    show_columns = list_columns",
        f"    edit_columns = list_columns",
        f"    add_columns = list_columns",
        f"    list_widget = BeautifulListWidget",
        f"    edit_widget = BeautifulFormWidget",
        f"    add_widget = BeautifulFormWidget",
        f"    show_widget = BeautifulFormWidget",
        "",
        f"    # Field sets for add and edit forms",
        f"    field_sets = {field_sets}",
        "",
        f"    # Enhanced search functionality",
        f"    search_columns = {list_columns}",
        "",
        f"    # Improved labels and descriptions",
        f"    label_columns = {{",
        f"        {', '.join([f"'{col}': '{col.replace('_', ' ').title()}'" for col in list_columns])}",
        f"    }}",
        f"    description_columns = {{",
        f"        {', '.join([f"'{col}': 'Enter the {col.replace('_', ' ')} here'" for col in list_columns])}",
        f"    }}",
        "",
        f"    # Custom formatters for better data presentation",
        f"    formatters_columns = {{",
        f"        'created_at': lambda x: x.strftime('%Y-%m-%d %H:%M:%S') if x else '',",
        f"        'updated_at': lambda x: x.strftime('%Y-%m-%d %H:%M:%S') if x else '',",
        f"    }}",
        "",
        f"    # Enhanced form field widgets",
        f"    add_form_extra_fields = {{",
        f"        {', '.join([generate_form_field(c) for c in table.columns if c.name in list_columns])}",
        f"    }}",
        "",
        f"    # Enable in-place editing",
        f"    can_edit = True",
        "",
        f"    # Custom actions",
        f"    @action('delete_all', 'Delete All', 'Are you sure you want to delete all records?', 'fa-trash', multiple=True)",
        f"    def delete_all(self, items):",
        f"        if isinstance(items, list):",
        f"            self.datamodel.delete_all(items)",
        f"            flash(f'Deleted {{len(items)}} records', 'success')",
        f"        else:",
        f"            flash('No records selected', 'warning')",
        f"        return redirect(request.referrer)",
        "",
        f"    @action('print', 'Print', 'Print the selected items?', 'fa-print', single=False)",
        f"    def print_items(self, items):",
        f"        if isinstance(items, list):",
        f"            return render_template('print_items.html', items=items, columns=self.list_columns)",
        f"        else:",
        f"            flash('No items selected', 'warning')",
        f"        return redirect(request.referrer)",
        "",
        f"    @action('export_csv', 'Export CSV', 'Export selected items to CSV?', 'fa-file-excel-o', single=False)",
        f"    def export_csv(self, items):",
        f"        if isinstance(items, list):",
        f"            csv_data = self.datamodel.export_as_csv(items)",
        f"            response = make_response(csv_data)",
        f"            response.headers['Content-Disposition'] = f'attachment; filename=export.csv'",
        f"            response.headers['Content-Type'] = 'text/csv'",
        f"            return response",
        f"        else:",
        f"            flash('No items selected', 'warning')",
        f"        return redirect(request.referrer)",
        "",
       f"    @action('bookmark', 'Bookmark', 'Bookmark selected items?', 'fa-bookmark', single=False)",
        f"    def bookmark_items(self, items):",
        f"        if isinstance(items, list):",
        f"            for item in items:",
        f"                item.is_bookmarked = True",
        f"            self.datamodel.bulk_update(items)",
        f"            flash(f'Bookmarked {{len(items)}} items', 'success')",
        f"        else:",
        f"            flash('No items selected', 'warning')",
        f"        return redirect(request.referrer)",
        "",
        f"    @action('merge', 'Merge', 'Merge selected items?', 'fa-compress', single=False)",
        f"    def merge_items(self, items):",
        f"        if isinstance(items, list) and len(items) > 1:",
        f"            # Implement merge logic here",
        f"            flash(f'Merged {{len(items)}} items', 'success')",
        f"        else:",
        f"            flash('Select at least two items to merge', 'warning')",
        f"        return redirect(request.referrer)",
        "",
        f"    @action('split', 'Split', 'Split selected item?', 'fa-scissors', single=True)",
        f"    def split_item(self, item):",
        f"        # Implement split logic here",
        f"        flash(f'Split item {{item}}', 'success')",
        f"        return redirect(request.referrer)",
        "",
        f"    @action('clone', 'Clone', 'Clone selected item?', 'fa-clone', single=True)",
        f"    def clone_item(self, item):",
        f"        new_item = self.datamodel.obj()",
        f"        for col in self.list_columns:",
        f"            setattr(new_item, col, getattr(item, col))",
        f"        self.datamodel.add(new_item)",
        f"        flash(f'Cloned item {{item}}', 'success')",
        f"        return redirect(request.referrer)",
        "",
        f"    @action('archive', 'Archive', 'Archive selected items?', 'fa-archive', single=False)",
        f"    def archive_items(self, items):",
        f"        if isinstance(items, list):",
        f"            for item in items:",
        f"                item.is_archived = True",
        f"            self.datamodel.bulk_update(items)",
        f"            flash(f'Archived {{len(items)}} items', 'success')",
        f"        else:",
        f"            flash('No items selected', 'warning')",
        f"        return redirect(request.referrer)",
        "",
        f"    @action('restore', 'Restore', 'Restore selected items?', 'fa-undo', single=False)",
        f"    def restore_items(self, items):",
        f"        if isinstance(items, list):",
        f"            for item in items:",
        f"                item.is_archived = False",
        f"            self.datamodel.bulk_update(items)",
        f"            flash(f'Restored {{len(items)}} items', 'success')",
        f"        else:",
        f"            flash('No items selected', 'warning')",
        f"        return redirect(request.referrer)",
        "",
        f"    @action('bulk_edit', 'Bulk Edit', 'Edit selected items?', 'fa-edit', single=False)",
        f"    def bulk_edit(self, items):",
        f"        if isinstance(items, list):",
        f"            return redirect(url_for('.bulk_edit_form', ids=','.join([str(item.id) for item in items])))",
        f"        else:",
        f"            flash('No items selected', 'warning')",
        f"        return redirect(request.referrer)",
        "",
        f"    @expose('/bulk_edit_form/<ids>')",
        f"    @has_access",
        f"    def bulk_edit_form(self, ids):",
        f"        items = self.datamodel.get_list_by_ids(ids.split(','))",
        f"        form = self.add_form()",
        f"        if request.method == 'POST':",
        f"            form = self.add_form(request.form)",
        f"            if form.validate():",
        f"                for item in items:",
        f"                    form.populate_obj(item)",
        f"                self.datamodel.bulk_update(items)",
        f"                return redirect(self.get_redirect())",
        f"        return self.render_template('bulk_edit.html', form=form, items=items)",
        "",
        f"    # Advanced Filtering and Sorting",
        f"    base_filters = []",
        f"    base_order = []",
        "",
        f"    def pre_add(self, item):",
        f"        # Set created_at and updated_at if they exist",
        f"        if hasattr(item, 'created_at'):",
        f"            item.created_at = datetime.datetime.now()",
        f"        if hasattr(item, 'updated_at'):",
        f"            item.updated_at = datetime.datetime.now()",
        "",
        f"    def pre_update(self, item):",
        f"        # Update updated_at if it exists",
        f"        if hasattr(item, 'updated_at'):",
        f"            item.updated_at = datetime.datetime.now()",
        "",
        f"    # User-Specific Actions",
        f"    @expose('/favorite/<pk>')",
        f"    @has_access",
        f"    def favorite(self, pk):",
        f"        item = self.datamodel.get(pk)",
        f"        if item:",
        f"            current_user.favorites.append(item)",
        f"            db.session.commit()",
        f"            flash(f'Added {{item}} to favorites', 'success')",
        f"        return redirect(request.referrer)",
        "",
        f"    @expose('/watchlist/<pk>')",
        f"    @has_access",
        f"    def watchlist(self, pk):",
        f"        item = self.datamodel.get(pk)",
        f"        if item:",
        f"            current_user.watchlist.append(item)",
        f"            db.session.commit()",
        f"            flash(f'Added {{item}} to watchlist', 'success')",
        f"        return redirect(request.referrer)",
        "",
        f"    @expose('/personalize')",
        f"    @has_access",
        f"    def personalize(self):",
        f"        if request.method == 'POST':",
        f"            current_user.list_columns = request.form.getlist('columns')",
        f"            current_user.list_order = request.form.get('order')",
        f"            db.session.commit()",
        f"            flash('View settings updated', 'success')",
        f"        return self.render_template('personalize.html', columns=self.list_columns, current_columns=current_user.list_columns, current_order=current_user.list_order)",
        "",
        f"    # Integration with External Services",
        f"    def post_add(self, item):",
        f"        # Example: Send email notification",
        f"        send_email_notification(f'New {{self.__class__.__name__}} added', f'A new {{self.__class__.__name__}} has been added: {{item}}')",
        "",
        f"    def post_update(self, item):",
        f"        # Example: Update external API",
        f"        update_external_api(item)",
        "",
        f"    def post_delete(self, item):",
        f"        # Example: Log to external service",
        f"        log_to_external_service(f'{{self.__class__.__name__}} deleted: {{item}}')",
        "",
        f"    # Type-specific validators",
        f"    validators_columns = {{",
        f"        {generate_validators(table)}",
        f"    }}",
        "",
        f"    # Relationship query fields",
        f"    {generate_form_query_rel_fields(table)}",
        "",
        f"    # Custom __repr__ method",
        f"    {generate_repr_method(table)}",
        ""
    ]
    generated_code.append([view_name, "\n".join(view_code)])
    generated_views.append((view_name, "ModelView", table.name, icon))

def generate_form_query_rel_fields(table: Any) -> str:
    """Generate form_query_rel_fields for foreign key relationships."""
    form_query_rel_fields = {}
    for fk in table.foreign_keys:
        referred_table = fk.column.table
        form_query_rel_fields[fk.parent.name] = f"db.session.query({snake_to_pascal(referred_table.name)})"

    if form_query_rel_fields:
        return f"    form_query_rel_fields = {form_query_rel_fields}"
    return ""

def generate_form_fields(table: Any, metadata: Any) -> str:
    """Generate form fields with appropriate widgets and validations based on column types."""
    form_fields = []
    for column in table.columns:
        if column.name in ['id', 'created_at', 'updated_at', 'created_by', 'updated_by']:
            continue  # Skip special fields

        if isinstance(column.type, String):
            length_validator = f", validators.Length(max={column.type.length})" if column.type.length else ""
            if 'password' in column.name:
                form_fields.append(
                    f"    '{column.name}': StringField('{column.name.capitalize()}', widget=BS3PasswordFieldWidget(), validators=[validators.DataRequired(){length_validator}]),"
                )
            else:
                form_fields.append(
                    f"    '{column.name}': StringField('{column.name.capitalize()}', widget=BS3TextFieldWidget(), validators=[validators.DataRequired(){length_validator}]),"
                )
        elif isinstance(column.type, Text):
            form_fields.append(
                f"    '{column.name}' : TextAreaField('{column.name.capitalize()}', widget=BS3TextFieldWidget(), validators=[validators.DataRequired()]),"
            )
        elif isinstance(column.type, Boolean):
            form_fields.append(
                f"    '{column.name}' : BooleanField('{column.name.capitalize()}'),"
            )
        elif isinstance(column.type, Integer):
            form_fields.append(
                f"    '{column.name}' : IntegerField('{column.name.capitalize()}', validators=[validators.DataRequired()]),"
            )
        elif isinstance(column.type, Float):
            form_fields.append(
                f"    '{column.name}' : FloatField('{column.name.capitalize()}', validators=[validators.DataRequired()]),"
            )
        elif isinstance(column.type, Numeric):
            form_fields.append(
                f"    '{column.name}' : DecimalField('{column.name.capitalize()}', validators=[validators.DataRequired()]),"
            )
        elif isinstance(column.type, Date):
            form_fields.append(
                f"    '{column.name}' : DateField('{column.name.capitalize()}', widget=DatePickerWidget(), validators=[validators.DataRequired()]),"
            )
        elif isinstance(column.type, DateTime):
            form_fields.append(
                f"    '{column.name}' : DateTimeField('{column.name.capitalize()}', widget=DateTimePickerWidget(), validators=[validators.DataRequired()]),"
            )
        elif isinstance(column.type, Time):
            form_fields.append(
                f"    '{column.name}' : TimeField('{column.name.capitalize()}'),"
            )
        elif isinstance(column.type, Enum):
            enum_choices = [(choice, choice) for choice in column.type.enums]
            form_fields.append(
                f"    '{column.name}' : SelectField('{column.name.capitalize()}', choices={enum_choices}, validators=[validators.DataRequired()]),"
            )
        elif isinstance(column.type, ForeignKey):
            related_table = column.foreign_keys[0].column.table
            related_model = snake_to_pascal(related_table.name)
            relationship_field = determine_relationship_name([column.name], table.name, related_table.name)
            form_fields.append(
                            f"    '{relationship_field}': QuerySelectField('{relationship_field.capitalize()}', query_factory=lambda: db.session.query({related_model}), widget=Select2Widget(), allow_blank=True),"
            )

    # Handle many-to-many relationships
    for related_table_name, related_table in metadata.tables.items():
        if is_association_table(related_table) and table.name in [
            fk.column.table.name for fk in related_table.foreign_keys
        ]:
            other_table = [fk.column.table for fk in related_table.foreign_keys if fk.column.table.name != table.name][0]
            related_model = snake_to_pascal(other_table.name)
            field_name = determine_remote_relationship_name('many-to-many', table.name, other_table.name)
            form_fields.append(
                f"    '{field_name}': QuerySelectMultipleField('{field_name.capitalize()}', query_factory=lambda: db.session.query({related_model}), widget=Select2ManyWidget()),"
            )

    if not form_fields:
        return "    # No fields to generate"

    return "\n".join(form_fields)



def generate_validators(table: Any) -> str:
    """Generate form validators based on column constraints."""
    validators = []
    for column in table.columns:
        column_validator = get_validators(column)
        if column_validator:
            validators[column.name] = f"[{', '.join(column_validator)}]"

    if validators:
        return f"    validators_columns = {validators}"
    return ""


def generate_repr_method(table):
    """Generate __repr__ method for the view."""
    display_column = next((col.name for col in table.columns if col.name in ['name', 'title', 'label']), table.columns[0].name)
    return f"""
    def __repr__(self):
        return self.{display_column}
    """

def get_validators(column):
    column_validators = []
    if not column.nullable and not column.primary_key:
        column_validators.append('validators.DataRequired()')
    if column.unique:
        column_validators.append('validators.Unique()')
    if isinstance(column.type, String):
        if column.type.length:
            column_validators.append(f'validators.Length(max={column.type.length})')
        if 'email' in column.name.lower():
            column_validators.append('validators.Email()')
    if 'url' in column.name.lower():
        column_validators.append('validators.URL()')
    if isinstance(column.type, (Integer, Float)):
        column_validators.append('validators.NumberRange()')
    return column_validators

def generate_master_detail_views(table: sa.Table, metadata: sa.MetaData, p: inflect.engine) -> None:
    """
    Generate MasterDetailViews for tables with foreign key relationships.

    This function creates MasterDetailView classes for tables that have foreign key
    relationships, allowing for a hierarchical view of related data.

    Args:
        table (sa.Table): SQLAlchemy Table object
        metadata (sa.MetaData): SQLAlchemy MetaData object
        p (inflect.engine): Inflect engine for plural/singular conversions

    Returns:
        None
    """
    for fk in table.foreign_keys:
        parent_table = metadata.tables[fk.column.table.name]
        parent_class_name = get_class_name(parent_table.name, p)
        child_class_name = get_class_name(table.name, p)
        icon = get_view_icon(table.name, "MasterDetailView")
        view_name = f"{parent_class_name}{child_class_name}MasterDetailView"
        view_code = [
            f"class {view_name}(MasterDetailView):",
            f"    datamodel = SQLAInterface({parent_class_name})",
            f"    related_views = [{child_class_name}ModelView]",
            f"    list_widget = BeautifulListWidget",
            f"    edit_widget = BeautifulFormWidget",
            f"    show_widget = BeautifulFormWidget",
            f"",
            f"    # Customize the master view",
            f"    list_columns = {get_list_columns(parent_table.columns)}",
            f"    show_columns = list_columns",
            f"",
            f"    # Customize the detail view",
            f"    list_title = '{snake_to_words(parent_table.name)} with {snake_to_words(table.name)}'",
            f"    add_title = 'Add {snake_to_words(parent_table.name)}'",
            f"    edit_title = 'Edit {snake_to_words(parent_table.name)}'",
            f"",
            f"    # Add any additional customization here",
            f"",
            f"    {generate_repr_method(parent_table)}",
            f""
        ]
        generated_code.append([view_name, "\n".join(view_code)])
        generated_views.append((view_name, "MasterDetailView", f"{parent_table.name}_{table.name}", icon))

def generate_multiple_views(table: sa.Table, metadata: sa.MetaData, p: inflect.engine) -> None:
    """
    Generate MultipleViews for tables with multiple related tables.

    This function creates MultipleView classes for tables that have relationships
    with multiple other tables, providing a consolidated view of related data.

    Args:
        table (sa.Table): SQLAlchemy Table object
        metadata (sa.MetaData): SQLAlchemy MetaData object
        p (inflect.engine): Inflect engine for plural/singular conversions

    Returns:
        None
    """
    if len(table.foreign_keys) > 1:
        class_name = get_class_name(table.name, p)
        icon = get_view_icon(table.name, "MultipleView")
        view_name = f"{class_name}MultipleView"
        related_views = [f"{get_class_name(fk.column.table.name, p)}ModelView" for fk in table.foreign_keys]
        view_code = [
            f"class {view_name}(MultipleView):",
            f"    views = [{', '.join(related_views)}]",
            f"",
            f"    # Customize the multiple view",
            f"    list_title = '{snake_to_words(table.name)} Multiple View'",
            f"",
            f"    # Add any additional customization here",
            f"",
            f"    @expose('/custom_view')",
            f"    @has_access",
            f"    def custom_view(self):",
            f"        # Add custom view logic here",
            f"        return self.render_template('custom_multiple_view.html', views=self.views)",
            f"",
            f"    {generate_repr_method(table)}",
            f""
        ]
        generated_code.append([view_name, "\n".join(view_code)])
        generated_views.append((view_name, "MultipleView", table.name, icon))

def generate_wizard_view(table: sa.Table, p: inflect.engine) -> None:
    """
    Generate a WizardView for complex forms with multiple steps.

    This function creates a WizardView class for tables with many columns,
    breaking down the form into multiple steps for easier data entry.

    Args:
        table (sa.Table): SQLAlchemy Table object
        p (inflect.engine): Inflect engine for plural/singular conversions

    Returns:
        None
    """
    class_name = get_class_name(table.name, p)
    view_name = f"{class_name}WizardView"
    icon = get_view_icon(table.name, "WizardView")
    columns = [c.name for c in table.columns]
    total_steps = math.ceil(len(columns) / 5)

    wizard_code = [
        f"class {view_name}(WizardView):",
        f"    datamodel = SQLAInterface({class_name})",
        f"    total_steps = {total_steps}",
        f"",
        f"    def __init__(self):",
        f"        super().__init__()",
        f"        self._init_steps()",
        f""
    ]

    for step in range(1, total_steps + 1):
        start_idx = (step - 1) * 5
        end_idx = min(step * 5, len(columns))
        step_columns = columns[start_idx:end_idx]

        wizard_code.extend([
            f"    @expose('/step{step}', methods=['GET', 'POST'])",
            f"    def step{step}(self):",
            f"        form = DynamicForm()",
            f"        {generate_form_fields(step_columns, table)}",
            f"",
            f"        if form.validate_on_submit():",
            f"            session['step{step}_data'] = form.data",
            f"            self.mark_step_complete({step})",
            f"            next_step = {step + 1 if step < total_steps else 1}",
            f"            return redirect(url_for(f'.step{{next_step}}'))",
            f"",
           f"        form_data = session.get('step{step}_data', {{}})",
        f"        form = DynamicForm(**form_data)",
        f"        widget = BeautifulFormWidget()",
        f"        return self.render_wizard({step}, widget(form), 'Step {step}: {', '.join(step_columns)}')",
        f""
    ]

    wizard_code.extend([
        f"    @expose('/submit', methods=['GET', 'POST'])",
        f"    def submit(self):",
        f"        if all(self.is_step_complete(step) for step in range(1, self.total_steps + 1)):",
        f"            # Combine data from all steps",
        f"            combined_data = {{}}",
        f"            for step in range(1, self.total_steps + 1):",
        f"                combined_data.update(session.get(f'step{{step}}_data', {{}}))",
        f"",
        f"            # Create new record",
        f"            item = self.datamodel.obj()",
        f"            for key, value in combined_data.items():",
        f"                setattr(item, key, value)",
        f"            self.datamodel.add(item)",
        f"",
        f"            # Clear session data",
        f"            for step in range(1, self.total_steps + 1):",
        f"                session.pop(f'step{{step}}_data', None)",
        f"                session.pop(f'{{self.__class__.__name__}}_step{{step}}_complete', None)",
        f"",
        f"            flash('Form submitted successfully!', 'success')",
        f"            return redirect(url_for('.step1'))",
        f"        else:",
        f"            flash('Please complete all steps before submitting.', 'warning')",
        f"            return redirect(url_for('.step1'))",
        f"",
        f"    {generate_repr_method(table)}",
        f""
    ])

    generated_code.append([view_name, "\n".join(wizard_code)])
    generated_views.append((view_name, "WizardView", table.name, icon))

def generate_graphql(table: sa.Table, p: inflect.engine) -> None:
    """
    Generate GraphQL schema and queries for a table.

    This function creates GraphQL ObjectTypes and queries for the given table,
    enabling GraphQL API access to the data.

    Args:
        table (sa.Table): SQLAlchemy Table object
        p (inflect.engine): Inflect engine for plural/singular conversions

    Returns:
        None
    """
    class_name = get_class_name(table.name, p)
    generated_code.append([
        f"{class_name}GraphQL",
        f"""
class {class_name}Object(SQLAlchemyObjectType):
    class Meta:
        model = {class_name}
        interfaces = (graphene.relay.Node, )

class {class_name}Connection(graphene.relay.Connection):
    class Meta:
        node = {class_name}Object

class Query(graphene.ObjectType):
    node = graphene.relay.Node.Field()
    all_{table.name} = SQLAlchemyConnectionField({class_name}Connection)

    {class_name.lower()} = graphene.Field({class_name}Object, id=graphene.Int())
    def resolve_{class_name.lower()}(self, info, id):
        return {class_name}.query.get(id)

schema.query = Query
"""
    ])

def generate_ModelRestApi(table: sa.Table, p: inflect.engine) -> None:
    """
    Generate a ModelRestApi for RESTful API access to a table.

    This function creates a ModelRestApi class for the given table, providing
    a RESTful API interface for CRUD operations on the data.

    Args:
        table (sa.Table): SQLAlchemy Table object
        p (inflect.engine): Inflect engine for plural/singular conversions

    Returns:
        None
    """
    class_name = get_class_name(table.name, p)
    view_name = f"{class_name}RestApi"
    columns = [c.name for c in table.columns]
    icon = get_view_icon(table.name, "ModelRestApiView")
    generated_code.append([
        view_name,
        f"""
class {view_name}(ModelRestApi):
    resource_name = '{table.name}'
    datamodel = SQLAInterface({class_name})
    allow_browser_login = True
    list_columns = {columns}
    show_columns = list_columns
    edit_columns = list_columns
    add_columns = list_columns

    # Customize API endpoints
    @expose('/custom_endpoint', methods=['GET'])
    def custom_endpoint(self):
        # Add custom API logic here
        return self.response(200, result='Custom endpoint reached')

    # Add data validation
    def pre_add(self, item):
        # Perform data validation before adding
        self.validate_data(item)
        super().pre_add(item)

    def pre_update(self, item):
        # Perform data validation before updating
        self.validate_data(item)
        super().pre_update(item)

    def validate_data(self, item):
        # Add custom validation logic here
        pass

    {generate_repr_method(table)}
"""
    ])
    generated_views.append((view_name, "ModelRestApi", table.name, icon))

def generate_chart_view(table: sa.Table, p: inflect.engine) -> None:
    """
    Generate a ChartView for data visualization of a table.

    This function creates a ChartView class for the given table, providing
    a visual representation of the data using charts.

    Args:
        table (sa.Table): SQLAlchemy Table object
        p (inflect.engine): Inflect engine for plural/singular conversions

    Returns:
        None
    """
    class_name = get_class_name(table.name, p)
    view_name = f"{class_name}ChartView"
    columns = [c.name for c in table.columns]
    icon = get_view_icon(table.name, "ChartView")
    chart_type = determine_chart_type(table)
    x_axis = get_suitable_x_axis(table)
    y_axis = get_suitable_y_axis(table)

    generated_code.append([
        view_name,
        f"""
class {view_name}(GroupByChartView):
    datamodel = SQLAInterface({class_name})
    chart_title = '{class_name} Distribution'
    label_columns = {class_name}ModelView.label_columns
    chart_type = '{chart_type}'
    definitions = [
        {{
            'label': '{y_axis} by {x_axis}',
            'group': '{x_axis}',
            'series': ['{y_axis}']
        }}
    ]

    # Add custom chart options
    chart_options = {{
        'legend': {{'position': 'bottom'}},
        'animation': {{'duration': 1000, 'easing': 'easeOutQuad'}}
    }}

    # Customize data query
    def query_init(self, query):
        # Add any custom query logic here
        return query

    # Add drill-down capability
    @expose('/drilldown/<col>/<val>')
    @has_access
    def drilldown(self, col, val):
        # Implement drill-down logic here
        pass

    {generate_repr_method(table)}
"""
    ])
    generated_views.append((view_name, "ChartView", table.name, icon))

def generate_form_field(column: sa.Column) -> str:
    """
    Generate a form field based on the column type.

    This function creates an appropriate form field for a given database column,
    taking into account the column's data type and constraints.

    Args:
        column (sa.Column): SQLAlchemy Column object

    Returns:
        str: String representation of the form field
    """
    if isinstance(column.type, sa.String):
        return f"'{column.name}': StringField('{column.name.replace('_', ' ').title()}', widget=BS3TextFieldWidget())"
    elif isinstance(column.type, sa.Integer):
        return f"'{column.name}': IntegerField('{column.name.replace('_', ' ').title()}')"
    elif isinstance(column.type, sa.Date):
        return f"'{column.name}': DateField('{column.name.replace('_', ' ').title()}', widget=DatePickerWidget())"
    elif isinstance(column.type, sa.Enum):
        choices = [(choice, choice) for choice in column.type.enums]
        return f"'{column.name}': SelectField('{column.name.replace('_', ' ').title()}', choices={choices}, widget=Select2Widget())"
    else:
        return f"'{column.name}': StringField('{column.name.replace('_', ' ').title()}', widget=BS3TextFieldWidget())"

def generate_validators(table: sa.Table) -> str:
    """
    Generate validators for table columns.

    This function creates appropriate validators for each column in the table,
    based on the column's constraints and data type.

    Args:
        table (sa.Table): SQLAlchemy Table object

    Returns:
        str: String representation of the validators dictionary
    """
    validators = []
    for column in table.columns:
        column_validators = []
        if not column.nullable and not column.primary_key:
            column_validators.append('validators.DataRequired()')
        if isinstance(column.type, sa.String) and column.type.length:
            column_validators.append(f'validators.Length(max={column.type.length})')
        if column_validators:
            validators.append(f"'{column.name}': [{', '.join(column_validators)}]")
    return ", ".join(validators)

def generate_form_query_rel_fields(table: sa.Table) -> str:
    """
    Generate form query relation fields for foreign key relationships.

    This function creates query fields for foreign key relationships in the table,
    enabling proper form handling for related data.

    Args:
        table (sa.Table): SQLAlchemy Table object

    Returns:
        str: String representation of the form_query_rel_fields dictionary
    """
    form_query_rel_fields = {}
    for fk in table.foreign_keys:
        referred_table = fk.column.table
        form_query_rel_fields[fk.parent.name] = f"db.session.query({get_class_name(referred_table.name, p)})"

    if form_query_rel_fields:
        return f"form_query_rel_fields = {form_query_rel_fields}"
    return ""

def generate_repr_method(table: sa.Table) -> str:
    """
    Generate a __repr__ method for a view class.

    This function creates a string representation method for the view class,
    typically using a primary identifier or name field of the table.

    Args:
        table (sa.Table): SQLAlchemy Table object

    Returns:
        str: String representation of the __repr__ method
    """
    display_column = next((col.name for col in table.columns if col.name in ['name', 'title', 'label']), table.columns[0].name)
    return f"""
    def __repr__(self):
        return str(self.{display_column})
    """

def determine_chart_type(table: sa.Table) -> str:
    """
    Determine the most suitable chart type for a table.

    This function analyzes the table structure to suggest an appropriate chart type
    for data visualization.

    Args:
        table (sa.Table): SQLAlchemy Table object

    Returns:
        str: Suggested chart type
    """
    date_columns = [col for col in table.columns if isinstance(col.type, (sa.Date, sa.DateTime))]
    if date_columns:
        return 'LineChart'
    else:
        return 'BarChart'

def get_suitable_x_axis(table: sa.Table) -> str:
    """
    Determine a suitable x-axis for a chart based on the table structure.

    This function suggests an appropriate column to use as the x-axis in a chart,
    preferring date/time columns or string columns if available.

    Args:
        table (sa.Table): SQLAlchemy Table object

    Returns:
        str: Name of the suggested x-axis column
    """
    date_columns = [col.name for col in table.columns if isinstance(col.type, (sa.Date, sa.DateTime))]
    if date_columns:
        return date_columns[0]
    else:
        string_columns = [col.name for col in table.columns if isinstance(col.type, (sa.String, sa.Text))]
        return string_columns[0] if string_columns else table.columns[0].name

def get_suitable_y_axis(table: sa.Table) -> str:
    """
    Determine a suitable y-axis for a chart based on the table structure.

    This function suggests an appropriate column to use as the y-axis in a chart,
    preferring numeric columns if available.

    Args:
        table (sa.Table): SQLAlchemy Table object

    Returns:
        str: Name of the suggested y-axis column
    """
    numeric_columns = [col.name for col in table.columns if isinstance(col.type, (sa.Integer, sa.Float, sa.Numeric))]
    return numeric_columns[0] if numeric_columns else table.columns[0].name

def generate_form_fields(columns: List[str], table: sa.Table) -> str:
    """
    Generate form fields for a set of columns.

    This function creates form fields for the given columns, using appropriate
    field types based on the column's data type.

    Args:
        columns (List[str]): List of column names
        table (sa.Table): SQLAlchemy Table object

    Returns:
        str: String representation of form field definitions
    """
    form_fields = []
    for column in columns:
        col_obj = table.columns[column]
        form_fields.append(f"form.{column} = {generate_form_field(col_obj)}")
    return "\n        ".join(form_fields)

def write_to_file(output_file: str) -> None:
    """
    Write the generated code to a file.

    This function writes all the generated view code to the specified output file.

    Args:
        output_file (str): Path to the output file

    Returns:
        None
    """
    with open(output_file, 'w') as f:
        for section_name, code in generated_code:
            f.write(f"# {section_name}\n")
            f.write(code)
            f.write("\n\n")

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Generate Flask-AppBuilder views from a database schema.")
    parser.add_argument("--uri", help="Database URI to connect to")
    parser.add_argument("--output", help="Output file to write the generated views")
    args = parser.parse_args()

    generate_views(args.uri)
    write_to_file(args.output)
    print(f"{len(generated_views)} Views have been generated successfully and written to {args.output}")
